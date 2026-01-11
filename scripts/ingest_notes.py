import sys
import os
import argparse

# Add backend to sys.path
sys.path.append(os.path.join(os.path.dirname(__file__), 'backend'))

from app.vectorstore.chroma_client import get_chroma_client
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter

def ingest_path(path: str):
    if not os.path.exists(path):
        print(f"Error: Path {path} does not exist.")
        return

    if os.path.isdir(path):
        print(f"Processing directory: {path}...")
        for root, dirs, files in os.walk(path):
            for file in files:
                file_path = os.path.join(root, file)
                ingest_file(file_path)
    else:
        ingest_file(path)

def ingest_file(file_path: str):
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in ['.pdf', '.docx', '.doc', '.txt', '.md']:
        return

    print(f"Loading {file_path}...")
    
    try:
        if ext == '.pdf':
            loader = PyPDFLoader(file_path)
            docs = loader.load()
        elif ext in ['.docx', '.doc']:
            try:
                from docx import Document as DocxDocument
                from langchain_core.documents import Document
                docx_doc = DocxDocument(file_path)
                text = "\n".join([p.text for p in docx_doc.paragraphs])
                docs = [Document(page_content=text, metadata={"source": file_path})]
            except ImportError:
                print("python-docx not installed. Cannot process DOCX.")
                return
        elif ext in ['.txt', '.md']:
            loader = TextLoader(file_path, encoding='utf-8')
            docs = loader.load()
        else:
            return
    except Exception as e:
        print(f"Error loading {file_path}: {e}")
        return

    if not docs:
        print(f"No content extracted from {file_path}")
        return

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    splits = text_splitter.split_documents(docs)

    print(f"Adding {len(splits)} chunks to vectorstore...")
    vectorstore = get_chroma_client()
    vectorstore.add_documents(splits)
    print(f"Successfully ingested {file_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ingest study materials into the vectorstore.")
    parser.add_argument("path", type=str, help="Path to a file or directory containing study materials.")
    args = parser.parse_args()
    
    ingest_path(args.path)
